import json
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from django.core.files.storage import FileSystemStorage
from .forms import ExcelUploadForm, WordUploadForm
from tools.libs.utils import excel_to_json, word_to_json
from tools.libs.txtToJson import txt_to_json, extract_code_name
import zipfile
import os
from django.core.files.storage import default_storage
from io import BytesIO, StringIO
from django.utils import timezone
import pandas as pd
from docx import Document

def excel_to_json_view(request):
    # Check if the request method is POST
    if request.method == 'POST':
        form = ExcelUploadForm(request.POST, request.FILES)
        if form.is_valid():
            json_files = []  # List to store JSON files for each sheet

            # Loop through each uploaded Excel file
            for excel_file in form.cleaned_data['files']:  # Using cleaned_data
                try:
                    # Read data from the Excel file, getting all sheets
                    excel_data = pd.read_excel(excel_file, sheet_name=None)  # Read all sheets from the Excel file
                    print(f"EXCEL_DATA IS :{excel_data}")

                    # Loop through each sheet in the Excel data
                    for sheet_name, df in excel_data.items():
                        # Convert the sheet to JSON format using the excel_to_json function
                        json_output = excel_to_json(df)  # excel_to_json() expects a DataFrame
                        print(json_output)

                        # Create a separate JSON filename for each sheet
                        json_filename = f"{excel_file.name.split('.')[0]}_{sheet_name}.json"
                        json_files.append((json_filename, json_output))  # Append each sheet's JSON to the list

                except Exception as e:
                    # Print an error message if there's an issue processing the Excel file
                    print(f"Error processing Excel file '{excel_file.name}': {e}")

            # Create a ZIP file to contain all the JSON files
            file_name_without_extension = os.path.splitext(excel_file.name)[0]
            zip_filename = file_name_without_extension + '_converted.zip' 
            response = HttpResponse(content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'

            # Write each JSON file to the ZIP archive
            with zipfile.ZipFile(response, 'w') as zip_file:
                for json_filename, json_string in json_files:
                    # Add each JSON file for each sheet to the ZIP
                    zip_file.writestr(json_filename, json_string)

            return response  # Return the response with the ZIP file

    else:
        # If the request method is not POST, create a new form
        form = ExcelUploadForm()

    # Render the form in the specified template
    return render(request, 'tool_excel_to_json.html', {'form': form})

def read_docx_content(file):
    document = Document(file)
    content = []

    for paragraph in document.paragraphs:
        content.append(paragraph.text.strip())
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                content.append(cell.text.strip())

    return "\n".join(content)

def word_to_json_view(request):
    if request.method == 'POST':
        form = WordUploadForm(request.POST, request.FILES)
        if form.is_valid():
            json_files = []
            for uploaded_file in form.cleaned_data['files']:
                try:
                    if uploaded_file.name.endswith('.docx'):
                        content = read_docx_content(uploaded_file)
                        json_output = word_to_json(content)
                        
                        # Kiểm tra nếu json_output là cấu trúc rỗng
                        if json_output == '{"mc_questions": []}':
                            print("Không tìm thấy câu hỏi hợp lệ.")
                        else:
                            print(json_output)

                        json_filename = f"{os.path.splitext(uploaded_file.name)[0]}.json"
                        json_files.append((json_filename, json_output))
                
                except Exception as e:
                    print(f"Lỗi khi xử lý tệp '{uploaded_file.name}': {e}")
                    # Nếu có lỗi, trả về cấu trúc JSON rỗng
                    json_filename = f"{os.path.splitext(uploaded_file.name)[0]}.json"
                    json_files.append((json_filename, json.dumps({"mc_questions": []}, ensure_ascii=False)))

            # Tạo file zip chứa các file JSON
            response = HttpResponse(content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename="converted_files.zip"'
            with zipfile.ZipFile(response, 'w') as zip_file:
                for json_filename, json_string in json_files:
                    zip_file.writestr(json_filename, json_string)
            return response
    else:
        form = WordUploadForm()

    return render(request, 'tool_word_to_json.html', {'form': form})

# def read_docx_content(file):
#     """Đọc nội dung từ tệp .docx."""
#     document = Document(file)  # Mở tệp .docx
#     content = []  # Khởi tạo danh sách để lưu nội dung

#     # Lặp qua tất cả các đoạn văn và bảng trong tài liệu
#     for paragraph in document.paragraphs:
#         content.append(paragraph.text.strip())  # Thêm nội dung đoạn văn vào danh sách
    
#     # Lặp qua tất cả các bảng trong tài liệu
#     for table in document.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 content.append(cell.text.strip())  # Thêm nội dung của ô vào danh sách

#     # Kết hợp tất cả nội dung thành một chuỗi, mỗi đoạn cách nhau bởi dấu xuống dòng
#     return "\n".join(content)

# def word_to_json_view(request):
#     if request.method == 'POST':
#         form = WordUploadForm(request.POST, request.FILES)
#         if form.is_valid():
#             json_files = []
#             for uploaded_file in form.cleaned_data['files']:
#                 try:
#                     # Đọc nội dung file .docx
#                     if uploaded_file.name.endswith('.docx'):
#                         content = read_docx_content(uploaded_file)
                    
#                     # Gọi hàm word_to_json để chuyển đổi nội dung sang JSON
#                     json_output = word_to_json(content)
#                     print(json_output)
#                     json_filename = f"{os.path.splitext(uploaded_file.name)[0]}.json"
#                     json_files.append((json_filename, json_output))
                
#                 except Exception as e:
#                     print(f"Lỗi khi xử lý tệp '{uploaded_file.name}': {e}")

#             # Tạo file zip chứa các file JSON
#             response = HttpResponse(content_type='application/zip')
#             response['Content-Disposition'] = f'attachment; filename="converted_files.zip"'
#             with zipfile.ZipFile(response, 'w') as zip_file:
#                 for json_filename, json_string in json_files:
#                     zip_file.writestr(json_filename, json_string)
#             return response
#     else:
#         form = WordUploadForm()

#     return render(request, 'tool_word_to_json.html', {'form': form})

def view_tools(request):
    return render(request,'view_tools.html')


def download_zip_file(request):
    # Generate or retrieve the zip file and return it as a response
    file_path = 'path/to/your/zip_file.zip'  # Replace with actual path or logic to create zip
    if default_storage.exists(file_path):
        with default_storage.open(file_path, 'rb') as zip_file:
            response = HttpResponse(zip_file.read(), content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename={file_path.split("/")[-1]}'
            return response
    return HttpResponse(status=404)

def txt_to_json_view(request):
    """Chuyển đổi văn bản từ textarea sang JSON và trả về file ZIP để tải xuống ngay lập tức."""
    if request.method == 'POST':
        texts = request.POST.getlist('texts')  # Lấy danh sách các chuỗi từ textarea

        # Tạo file ZIP trong bộ nhớ
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for idx, text_content in enumerate(texts):
                if text_content.strip():
                    # Lấy tên tệp từ nội dung hoặc đặt tên mặc định
                    file_name = extract_code_name(text_content) or f'text_{idx + 1}'
                    file_like = StringIO(text_content)

                    # Chuyển đổi nội dung thành JSON
                    json_output = txt_to_json(file_like, file_name)
                    json_file_name = f'{file_name}.json'

                    # Lưu dữ liệu JSON vào file ZIP
                    zip_file.writestr(json_file_name, json_output)

        # Đặt con trỏ của zip_buffer về đầu để đọc
        zip_buffer.seek(0)

        # Đặt tên file ZIP theo timestamp
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        zip_filename = f'json_files_{timestamp}.zip'

        # Tạo response để tải xuống file ZIP
        response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'

        return response

    return render(request, 'tool_txt_to_json.html')






